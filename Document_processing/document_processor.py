import os
from docx import Document
from settings import output_directory


class DocumentProcessor:
    """
    класс DocumentProcessor позволяет заполнять шаблоны документов данными
     из dataset и сохранять полученные заполненные документы
    """
    def __init__(self):
        self.templates = {}

    def load_template(self, template_name, order):
        self.templates[template_name] = order

    def process_dataset(self, dataset):
        for template_name, order in sorted(self.templates.items(), key=lambda x: x[1]):
            if template_name in dataset:
                template_data = dataset[template_name]

                if self.fill_template(template_name, template_data):
                    print(f"Document generated successfully based on {template_name}.")
                else:
                    print(f"Error: Failed to generate document based on {template_name}.")
            else:
                print(f"Error: Template {template_name} not found in dataset.")

    def fill_template(self, template_name, data):
        try:
            doc = Document(template_name)
            # Замена шаблонов в параграфе
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    if f'${key}' in paragraph.text:
                        for run in paragraph.runs:
                            if f'${key}' in run.text:
                                run.text = run.text.replace(f'${key}', value)
                                break
            # Замена шаблонов в таблице
            for table in doc.tables:
                total_payment_sum = 0
                row_index = 1
                for data_row in data.get("rows", []):
                    if row_index < len(table.rows):
                        row = table.rows[row_index]
                        for cell_index, cell in enumerate(row.cells):
                            if cell.text == "$pp_number":
                                cell.text = str(data_row.get("pp_number", ""))
                            elif cell.text == "$lot_number":
                                cell.text = str(data_row.get("lot_number", ""))
                            elif cell_index == 3 and "pp_date" in data_row:
                                cell.text = data_row.get("pp_date", "")
                            elif cell_index == 4:
                                cell.text = data_row.get("winer_name", "")
                            elif cell_index == 5:
                                cell.text = data_row.get("payment_purpose", "")
                            elif cell_index == 6:
                                cell.text = str(data_row.get("payment_sum", 0))
                                total_payment_sum += data_row.get("payment_sum", 0)
                            # elif cell.text == "$total":
                            #     cell.text = str(total_payment_sum)

                        row_index += 1

                # Замена ключа "$total" на значение total_payment_sum
                for row in table.rows:
                    for cell in row.cells:
                        if "$total" in cell.text:
                            cell.text = cell.text.replace("$total", str(total_payment_sum))

            output_filename = os.path.join(output_directory, f'output_{os.path.basename(template_name)}')
            doc.save(output_filename)

            return True
        except Exception as e:
            print(f"Error: {e}")
            return False

